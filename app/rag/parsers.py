"""文档解析器 — 从企业知识库 RAG 项目移植，支持更多格式。"""

from __future__ import annotations

from pathlib import Path

from langchain_core.documents import Document

from app.utils.logger import get_logger

logger = get_logger(__name__)


def parse_pdf(file_path: Path) -> list[Document]:
    """使用 PyMuPDF 解析 PDF，文本提取质量远优于 PyPDFLoader。"""
    import fitz  # pymupdf

    docs = []
    pdf = fitz.open(str(file_path))
    for page_num in range(len(pdf)):
        page = pdf[page_num]
        text = page.get_text("text")
        if text.strip():
            docs.append(Document(
                page_content=text,
                metadata={"source": str(file_path), "page": page_num},
            ))
    pdf.close()
    logger.info("PyMuPDF 解析 %s → %d 页", file_path.name, len(docs))
    return docs


def parse_txt(file_path: Path) -> list[Document]:
    """解析纯文本文件。"""
    text = file_path.read_text(encoding="utf-8")
    if not text.strip():
        return []
    return [Document(
        page_content=text,
        metadata={"source": str(file_path), "page": 0},
    )]


def parse_markdown(file_path: Path) -> list[Document]:
    """解析 Markdown 文件。"""
    return parse_txt(file_path)  # 同纯文本处理


def parse_docx(file_path: Path) -> list[Document]:
    """解析 Word 文档。"""
    from docx import Document as DocxDocument
    doc = DocxDocument(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    if not text.strip():
        return []
    return [Document(
        page_content=text,
        metadata={"source": str(file_path), "page": 0},
    )]


# 解析器注册表
PARSER_REGISTRY: dict[str, callable] = {
    ".pdf": parse_pdf,
    ".txt": parse_txt,
    ".md": parse_markdown,
    ".markdown": parse_markdown,
    ".docx": parse_docx,
    ".doc": parse_docx,
}


def parse_document(file_path: Path) -> list[Document]:
    """根据文件后缀选择解析器。"""
    suffix = file_path.suffix.lower()
    parser = PARSER_REGISTRY.get(suffix)
    if parser is None:
        raise ValueError(f"不支持的文件格式: {suffix}，支持: {list(PARSER_REGISTRY.keys())}")
    return parser(file_path)
